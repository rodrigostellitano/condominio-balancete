
from pydoc import doc
import pandas as pd
from unidecode import unidecode
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Pt, Inches

from repository.excel_repository import ExcelRepository
from services.finance_service import FinanceService
from utils.date_utils import Utils







def extract_house_debit(df):
    df_debito = df[["Casas com debito", "Mes", "Ano"]].dropna()

    MESES_ABREV = {
        1: "Jan", 2: "Fev", 3: "Mar", 4: "Abr",
        5: "Mai", 6: "Jun", 7: "Jul", 8: "Ago",
        9: "Set", 10: "Out", 11: "Nov", 12: "Dez"
    }

    total_casas_em_debito = len(df_debito["Casas com debito"])
    print("Total casas em debito:", total_casas_em_debito)
    resultado = []

    for casa, grupo in df_debito.groupby("Casas com debito"):
        anos = {}

        for _, row in grupo.iterrows():
            ano = int(row["Ano"])
            mes = int(row["Mes"])
            anos.setdefault(ano, []).append(mes)

        partes = []
        for ano, meses in sorted(anos.items()):
            meses_ordenados = sorted(meses)
            meses_txt = "/".join(MESES_ABREV[m] for m in meses_ordenados)
            partes.append(f"{meses_txt}/{str(ano)[-2:]}")

        resultado.append(f"Casa {int(casa):02d} ({' | '.join(partes)})")

    return resultado, total_casas_em_debito



def extract_names(df_aux):
    df_aux = df_aux[["Titulo", "Info"]].dropna()
    names = dict(zip(df_aux["Titulo"], df_aux["Info"]))
    sindico = names.get("Sindico", 0)
    president = names.get("Presidente", 0)
    cons_1 = names.get("Cons Fiscal 1", 0)
    cons_2 = names.get("Cons Fiscal 2", 0)
    
    
    return {
        "Síndico": sindico,
        "Presidente": president,
        "Fiscal 1": cons_1,
        "Fiscal 2": cons_2,
    }

def extract_house_payed(df):
    df_pago = df[["Cotas Atrasadas Pagas", "Mes Atrasado", "Ano Atrasado"]].dropna()

    MESES_ABREV = {
        1: "Jan", 2: "Fev", 3: "Mar", 4: "Abr",
        5: "Mai", 6: "Jun", 7: "Jul", 8: "Ago",
        9: "Set", 10: "Out", 11: "Nov", 12: "Dez"
    }

    resultado = []

    for casa, grupo in df_pago.groupby("Cotas Atrasadas Pagas"):
        anos = {}

        for _, row in grupo.iterrows():
            ano = int(row["Ano Atrasado"])
            mes = int(row["Mes Atrasado"])
            anos.setdefault(ano, []).append(mes)

        partes = []
        for ano, meses in sorted(anos.items()):
            meses_ordenados = sorted(meses)
            meses_txt = "/".join(MESES_ABREV[m] for m in meses_ordenados)
            partes.append(f"{meses_txt}/{str(ano)[-2:]}")

        resultado.append(f" {int(casa):02d}({' | '.join(partes)})")

    print(resultado)
    
    return resultado







def add_finance_line(doc, descricao, valor, tab_inches=6):
    """
    Adiciona uma linha com pontinhos até o valor.
    """
    # Cria parágrafo
    p = doc.add_paragraph()
    
    # Define tab stop com líder de pontinhos
    tab_stop = p.paragraph_format.tab_stops.add_tab_stop(
        Inches(tab_inches), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    
    # Adiciona run com texto + tab
    run = p.add_run(f"{descricao}\tR$ {valor:,.2f}")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    
    return p


def generate_report(
    date_info,
    previous_month_value,
    total_entrace,
    details_entrace,
    total_exit,
    details_exit,
    people,
    total_final,
    houses_debit,
    total_casas_em_debito,
    houses_paid
):
    doc = Document()

    # Título

    title = doc.add_heading(
        f"Balancete do Condomínio da Vila Santa Maria",
        level=1
    )
    title2 = doc.add_heading(f"Mês de {date_info['Current Month']} {date_info['Current Year']}", level=1)

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)  # tira espaço depois
    title.paragraph_format.space_before = Pt(0) 
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2.paragraph_format.space_after = Pt(12)
    title2.paragraph_format.space_before = Pt(0) # tira espaço depois

    # --- Datas ---
    add_finance_line(
        doc,
        f"Saldo do Balancete de {date_info['Previus Month']} {date_info['Previus Year']}",
        previous_month_value
    )
   
    # --- Valores ---
    receita = doc.add_heading(f"Receita de {date_info['Current Month']} {date_info['Current Year']}", level=2)

    for run in receita.runs:
        run.underline = True

    
    add_finance_line(
    doc,
    f"Taxa de Condomínio: {details_entrace['Qtd Cotas Pagas']}/30 cotas",
    details_entrace['Total Cotas']
)

    if details_entrace['Qtd Cotas Pagas Multa'] > 0:
        add_finance_line(
        doc, 
        f"Taxa de Condomínio Atrasada: {details_entrace['Qtd Cotas Pagas Multa']} cotas ",
        details_entrace['Total Cotas com Multa']
        
    )
        doc.add_paragraph(houses_paid)
   
    if details_entrace['Qtd Tag'] > 0:
        add_finance_line(
        doc, 
        f"Tag(s): {details_entrace['Qtd Tag']} tags",
        details_entrace['Total Tag']
        )
    
    # Passe apenas o número
    receita = add_finance_line(doc, "Total de Receitas", total_entrace)

    
    

    despesas = doc.add_heading(f"Despesas de {date_info['Current Month']} {date_info['Current Year']}", level=2)

    for run in despesas.runs:
        run.underline = True

    # --- Detalhes Saída ---
    
    for k, v in details_exit.items():
        
        add_finance_line(doc, k, v)
    
    add_finance_line(doc, "Total de Despesas", total_exit)

    add_finance_line(doc, f"Saldo do Balancete de {date_info['Current Month']}/{date_info['Current Year']}", total_final)
    doc.add_paragraph("- Comprovantes em poder do Síndico para conferência")
    doc.add_paragraph(f"- Encontra-se em aberto : {total_casas_em_debito} cotas")
    doc.add_paragraph("Relação das casas em débito:")
    for casa in houses_debit:
        doc.add_paragraph(f"- {casa}")

    # --- Responsáveis ---
    data = doc.add_heading(f"Rio de Janeiro,     de {date_info['Current Month']} de {date_info['Current Year']}", level=2)
    data.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cargo, nome in people.items():
        if "1" in cargo or "2" in cargo:
            doc.add_paragraph(f"{nome} - Fiscal: ________________________________________________ ")
        else:
            doc.add_paragraph(f"{nome} - {cargo}: ________________________________________________")

        

    

    # Salvar
    file_name = f"Balancete_{date_info['Current Month']}_{date_info['Current Year']}.docx"
    doc.save(file_name)

    return file_name






def main():
    #Load Excel Repository
    repository = ExcelRepository()
    df, df_aux = repository.load()

    finance = FinanceService()
    entrace = finance.calculate_entrace(df, df_aux)
    result_exit, df_saida = finance.calculate_exit(df)

    
    value_previus_month,current_month, total_entrace, df_entrada = entrace
    
    print(current_month)

    utility = Utils()
    
    date = utility.ajust_date(current_month)
    
    
    df_names = extract_names(df_aux)
    
    total = finance.calculate_total(value_previus_month,total_entrace,result_exit)
    
    df_debito, total_casas_em_debito= extract_house_debit(df)
    
    houses_paid =  extract_house_payed(df)

    print("Date:", date)
    
    print("Previous Month Entrace:", value_previus_month)

    print("Total Entrace:", total_entrace)

    print("Details Entrace:", df_entrada)
    print("Result Exit:", result_exit)
    print("Details Exit:", df_saida)
    print(df_names)
    print("Total Final:", total)
    print("Houses with debit:", df_debito)


    file = generate_report(
        date_info=date,
        previous_month_value=value_previus_month,
        total_entrace=total_entrace,
        details_entrace=df_entrada,
        total_exit=result_exit,
        details_exit=df_saida,
        houses_debit=df_debito,
        total_casas_em_debito=total_casas_em_debito,
        people=df_names,
        total_final=total,
        houses_paid=houses_paid
        
    )




if __name__ == "__main__":
    main()