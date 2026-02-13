
#reports/report_generator.py

from docx import Document
from docx.shared import Pt
from unidecode import unidecode
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Pt, Inches

class ReportGenerator:

    def add_finance_line(self,doc, descricao, valor, tab_inches=6):
        """
        Adiciona uma linha com pontinhos até o valor.
        """
        # Cria parágrafo
        p = doc.add_paragraph()
        
        # Define tab stop com líder de pontinhos
        tab_stop = p.paragraph_format.tab_stops.add_tab_stop(
            Inches(tab_inches), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        
        # Adiciona run com texto + tab
        run = p.add_run(f"{descricao}\tR$ {valor:,.2f}")
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        
        return p


    def generate_report(self,
        date_info,
        previous_month_value,
        total_entrace,
        details_entrace,
        total_exit,
        details_exit,
        people,
        total_final,
        houses_debit,
        total_casas_em_debito,
        houses_paid
    ):
        doc = Document()

        # Título

        title = doc.add_heading(
            f"Balancete do Condomínio da Vila Santa Maria",
            level=1
        )
        title2 = doc.add_heading(f"Mês de {date_info['Current Month']} {date_info['Current Year']}", level=1)

        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(0)  # tira espaço depois
        title.paragraph_format.space_before = Pt(0) 
        title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title2.paragraph_format.space_after = Pt(12)
        title2.paragraph_format.space_before = Pt(0) # tira espaço depois

        # --- Datas ---
        self.add_finance_line(
            doc,
            f"Saldo do Balancete de {date_info['Previus Month']} {date_info['Previus Year']}",
            previous_month_value
        )
    
        # --- Valores ---
        receita = doc.add_heading(f"Receita de {date_info['Current Month']} {date_info['Current Year']}", level=2)

        for run in receita.runs:
            run.underline = True

        
        self.add_finance_line(
        doc,
        f"Taxa de Condomínio: {details_entrace['Qtd Cotas Pagas']}/30 cotas",
        details_entrace['Total Cotas']
    )

        if details_entrace['Qtd Cotas Pagas Multa'] > 0:
            self.add_finance_line(
            doc, 
            f"Taxa de Condomínio Atrasada: {details_entrace['Qtd Cotas Pagas Multa']} cotas ",
            details_entrace['Total Cotas com Multa']
            
        )
            doc.add_paragraph(houses_paid)
    
        if details_entrace['Qtd Tag'] > 0:
            self.add_finance_line(
            doc, 
            f"Tag(s): {details_entrace['Qtd Tag']} tags",
            details_entrace['Total Tag']
            )
        
        # Passe apenas o número
        receita = self.add_finance_line(doc, "Total de Receitas", total_entrace)

        
        

        despesas = doc.add_heading(f"Despesas de {date_info['Current Month']} {date_info['Current Year']}", level=2)

        for run in despesas.runs:
            run.underline = True

        # --- Detalhes Saída ---
        
        for k, v in details_exit.items():
            
            self.add_finance_line(doc, k, v)
        
        self.add_finance_line(doc, "Total de Despesas", total_exit)

        self.add_finance_line(doc, f"Saldo do Balancete de {date_info['Current Month']}/{date_info['Current Year']}", total_final)
        doc.add_paragraph("- Comprovantes em poder do Síndico para conferência")
        doc.add_paragraph(f"- Encontra-se em aberto : {total_casas_em_debito} cotas")
        doc.add_paragraph("Relação das casas em débito:")
        for casa in houses_debit:
            doc.add_paragraph(f"- {casa}")

        # --- Responsáveis ---
        data = doc.add_heading(f"Rio de Janeiro,     de {date_info['Current Month']} de {date_info['Current Year']}", level=2)
        data.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for cargo, nome in people.items():
            if "1" in cargo or "2" in cargo:
                doc.add_paragraph(f"{nome} - Fiscal: ________________________________________________ ")
            else:
                doc.add_paragraph(f"{nome} - {cargo}: ________________________________________________")

            

        

        # Salvar
        file_name = f"Balancete_{date_info['Current Month']}_{date_info['Current Year']}.docx"
        doc.save(file_name)

        return file_name
